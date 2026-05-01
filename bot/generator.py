"""
Генерация пакета документов:
  • Счёт на оплату  (PDF  ← WeasyPrint ← Jinja2/HTML)
  • УПД Статус 1    (PDF  ← WeasyPrint ← Jinja2/HTML)
  • Договор поставки (DOCX ← docxtpl)
  • XML для ЭДО      (XML  ← lxml,  формат ON_NSCHFDOPPR v5.03)
"""
import io
import os
import uuid
from datetime import date, datetime
from decimal import Decimal, ROUND_HALF_UP
from typing import Any

import jinja2
from weasyprint import HTML as WeasyHTML
from docxtpl import DocxTemplate
from lxml import etree
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

import config
from num2words_ru import amount_to_words

# ─────────────────────── VAT helpers ──────────────────────────────

_TWO = Decimal('0.01')


def _r(x: Decimal) -> Decimal:
    return x.quantize(_TWO, rounding=ROUND_HALF_UP)


_MONTHS_GEN = [
    'января', 'февраля', 'марта', 'апреля', 'мая', 'июня',
    'июля', 'августа', 'сентября', 'октября', 'ноября', 'декабря',
]


def _date_ru(d: date) -> str:
    return f"{d.day} {_MONTHS_GEN[d.month - 1]} {d.year} г."


def _date_short(d: date) -> str:
    return d.strftime('%d.%m.%Y')


def enrich_items(raw_items: list[dict], delivery: Decimal) -> tuple[list[dict], Decimal, Decimal, Decimal]:
    """
    Добавляет к каждому товару: price_without_vat, total_without_vat, vat_amount, total_with_vat.
    Возвращает (enriched_items, total_excl_vat, total_vat, total_incl_vat).
    """
    result = []
    for it in raw_items:
        price_incl  = Decimal(str(it['price']))
        qty         = Decimal(str(it.get('qty', 1)))
        price_excl  = _r(price_incl / (1 + config.VAT_RATE))
        total_incl  = _r(price_incl * qty)
        total_excl  = _r(price_excl * qty)
        vat_amt     = _r(total_incl - total_excl)
        result.append({**it,
                       'price_without_vat': price_excl,
                       'total_without_vat': total_excl,
                       'vat_amount':        vat_amt,
                       'total_with_vat':    total_incl})

    if delivery and delivery > 0:
        d = Decimal(str(delivery))
        d_excl = _r(d / (1 + config.VAT_RATE))
        d_vat  = _r(d - d_excl)
        result.append({
            'name': 'Доставка', 'qty': '', 'unit': 'шт',
            'price': d, 'price_without_vat': d_excl,
            'total_without_vat': d_excl, 'vat_amount': d_vat, 'total_with_vat': d,
        })

    total_excl = _r(sum(i['total_without_vat'] for i in result))
    total_vat  = _r(sum(i['vat_amount']         for i in result))
    total_incl = _r(sum(i['total_with_vat']     for i in result))
    return result, total_excl, total_vat, total_incl


# ─────────────────────── Jinja2 env ───────────────────────────────

def _jinja() -> jinja2.Environment:
    env = jinja2.Environment(
        loader=jinja2.FileSystemLoader(config.TEMPLATES_DIR),
        autoescape=jinja2.select_autoescape(['html']),
    )
    env.filters['money'] = lambda v: f"{v:,.2f}".replace(',', ' ')
    return env


# ─────────────────────── Invoice PDF ──────────────────────────────

def generate_invoice_pdf(data: dict) -> bytes:
    items, total_excl, total_vat, total_incl = enrich_items(
        data['items'], Decimal(str(data.get('delivery', 0)))
    )
    doc_date     = data['date']
    payment_date = doc_date.replace(day=min(doc_date.day + config.PAYMENT_DAYS, 28))

    html = _jinja().get_template('invoice.html').render(
        doc_number    = data['doc_number'],
        doc_date      = _date_ru(doc_date),
        payment_date  = _date_short(payment_date),
        seller        = config,
        buyer         = data['buyer'],
        items         = items,
        total_excl    = total_excl,
        total_vat     = total_vat,
        total_incl    = total_incl,
        total_words   = amount_to_words(total_incl),
        basis         = data.get('basis', ''),
        item_count    = sum(1 for i in items if i.get('qty') != ''),
    )
    return WeasyHTML(string=html, base_url=config.TEMPLATES_DIR).write_pdf()


# ─────────────────────── UPD PDF ──────────────────────────────────

def generate_upd_pdf(data: dict) -> bytes:
    items, total_excl, total_vat, total_incl = enrich_items(
        data['items'], Decimal(str(data.get('delivery', 0)))
    )
    doc_date = data['date']

    html = _jinja().get_template('upd.html').render(
        doc_number   = data['doc_number'],
        doc_date     = _date_ru(doc_date),
        doc_date_s   = _date_short(doc_date),
        seller       = config,
        buyer        = data['buyer'],
        items        = items,
        total_excl   = total_excl,
        total_vat    = total_vat,
        total_incl   = total_incl,
        vat_label    = config.VAT_LABEL,
    )
    return WeasyHTML(string=html, base_url=config.TEMPLATES_DIR).write_pdf()


# ─────────────────────── Contract PDF ────────────────────────────

def generate_contract_pdf(data: dict) -> bytes:
    items, _, _, total_incl = enrich_items(
        data['items'], Decimal(str(data.get('delivery', 0)))
    )
    doc_date = data['date']
    buyer    = data['buyer']

    items_desc = '; '.join(
        f"{i['name']} ({i['qty']} {i['unit']}) — {i['total_with_vat']:.2f} руб."
        for i in items if i.get('qty') not in ('', None)
    )

    html = _jinja().get_template('contract.html').render(
        doc_number      = data['doc_number'],
        date_day        = doc_date.day,
        date_month      = _MONTHS_GEN[doc_date.month - 1],
        date_year       = doc_date.year,
        contract_end    = f"31 декабря {doc_date.year} г.",
        seller_city     = config.SELLER_CITY,
        buyer_name      = buyer.get('name', ''),
        buyer_inn       = buyer.get('inn', ''),
        buyer_kpp       = buyer.get('kpp', ''),
        buyer_address   = buyer.get('address', ''),
        buyer_director  = buyer.get('director', ''),
        buyer_rs        = buyer.get('rs', ''),
        buyer_bank_name = buyer.get('bank_name', ''),
        buyer_bik       = buyer.get('bik', ''),
        buyer_ks        = buyer.get('ks', ''),
        seller_name     = config.SELLER_NAME,
        seller_inn      = config.SELLER_INN,
        seller_address  = config.SELLER_ADDRESS,
        seller_rs       = config.SELLER_RS,
        seller_bank_name= config.SELLER_BANK_NAME,
        seller_bik      = config.SELLER_BIK,
        seller_ks       = config.SELLER_KS,
        seller_signature= config.SELLER_SIGNATURE,
        items_desc      = items_desc,
        total_amount    = f"{total_incl:.2f}",
        total_words     = amount_to_words(total_incl),
    )
    return WeasyHTML(string=html, base_url=config.TEMPLATES_DIR).write_pdf()


# ─────────────────────── Contract DOCX ────────────────────────────

def generate_contract_docx(data: dict) -> bytes:
    tmpl_path = os.path.join(config.TEMPLATES_DIR, 'contract.docx')
    if not os.path.exists(tmpl_path):
        _build_contract_template(tmpl_path)

    items, _, _, total_incl = enrich_items(
        data['items'], Decimal(str(data.get('delivery', 0)))
    )
    doc_date = data['date']
    buyer    = data['buyer']

    items_desc = '; '.join(
        f"{i['name']} ({i['qty']} {i['unit']}) — {i['total_with_vat']:.2f} руб."
        for i in items if i.get('qty') not in ('', None)
    )

    tpl = DocxTemplate(tmpl_path)
    tpl.render({
        'number':           data['doc_number'],
        'date_day':         doc_date.day,
        'date_month':       _MONTHS_GEN[doc_date.month - 1],
        'date_year':        doc_date.year,
        'buyer_name':       buyer.get('name', ''),
        'buyer_inn':        buyer.get('inn', ''),
        'buyer_kpp':        buyer.get('kpp', ''),
        'buyer_address':    buyer.get('address', ''),
        'buyer_director':   buyer.get('director', ''),
        'buyer_rs':         buyer.get('rs', ''),
        'buyer_bank':       buyer.get('bank_name', ''),
        'buyer_bik':        buyer.get('bik', ''),
        'buyer_ks':         buyer.get('ks', ''),
        'seller_name':      config.SELLER_NAME,
        'seller_inn':       config.SELLER_INN,
        'seller_address':   config.SELLER_ADDRESS,
        'seller_rs':        config.SELLER_RS,
        'seller_bank':      config.SELLER_BANK_NAME,
        'seller_bik':       config.SELLER_BIK,
        'seller_ks':        config.SELLER_KS,
        'seller_signature': config.SELLER_SIGNATURE,
        'items_desc':       items_desc,
        'total_amount':     f"{total_incl:.2f}",
        'total_words':      amount_to_words(total_incl),
        'contract_end':     f"31 декабря {doc_date.year} г.",
    })
    buf = io.BytesIO()
    tpl.save(buf)
    return buf.getvalue()


def _build_contract_template(path: str) -> None:
    """Создаёт шаблон договора с Jinja2-плейсхолдерами (docxtpl-совместимый)."""
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(1.5)

    # Заголовок
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('ДОГОВОР ПОСТАВКИ № {{number}}')
    r.bold = True; r.font.size = Pt(14)

    p2 = doc.add_paragraph()
    p2.add_run('г. Элиста')
    p2.add_run('\t\t\t\t\t\t\t\t')
    p2.add_run('«{{date_day}}» {{date_month}} {{date_year}} г.')

    doc.add_paragraph()

    # Преамбула
    pre = doc.add_paragraph()
    pre.add_run('{{buyer_name}}').bold = True
    pre.add_run(
        ', ИНН {{buyer_inn}}, КПП {{buyer_kpp}}, именуемое в дальнейшем «Покупатель», '
        'в лице {{buyer_director}}, действующего на основании Устава, с одной стороны, и '
    )
    pre.add_run(config.SELLER_NAME).bold = True
    pre.add_run(
        f', ИНН {config.SELLER_INN}, именуемый(ая) в дальнейшем «Поставщик», '
        'с другой стороны, заключили настоящий договор о нижеследующем:'
    )

    doc.add_paragraph()

    _sections = [
        ('1. ПРЕДМЕТ ДОГОВОРА', [
            '1.1. Поставщик обязуется поставить, а Покупатель — принять и оплатить товар: {{items_desc}}.',
            '1.2. Наименование, характеристики, количество и цена каждой партии товара определяются счётами, выставленными Поставщиком.',
            '1.3. Поставщик гарантирует, что поставляемый товар полностью ему принадлежит и не обременён правами третьих лиц.',
        ]),
        ('2. ЦЕНА И ПОРЯДОК ОПЛАТЫ', [
            '2.1. Стоимость товара по настоящему договору составляет {{total_amount}} руб. ({{total_words}}) с учётом НДС 5%.',
            '2.2. Покупатель оплачивает счёт в течение 3 (трёх) рабочих дней с даты его выставления.',
            '2.3. Датой оплаты считается день поступления денежных средств на расчётный счёт Поставщика.',
        ]),
        ('3. ПОСТАВКА И ПЕРЕХОД ПРАВА СОБСТВЕННОСТИ', [
            '3.1. Право собственности на товар переходит к Покупателю с момента его передачи, фиксируемой подписанием УПД.',
            '3.2. Товар отпускается по факту поступления оплаты на р/с Поставщика, самовывозом, при наличии доверенности и паспорта, если иное не оговорено дополнительно.',
        ]),
        ('4. КАЧЕСТВО ТОВАРА', [
            '4.1. Товар должен соответствовать качественным характеристикам, установленным для данного вида товара.',
            '4.2. Покупатель обязан осмотреть товар при получении и предъявить претензию по явным недостаткам в течение 3 (трёх) рабочих дней.',
        ]),
        ('5. ОТВЕТСТВЕННОСТЬ СТОРОН', [
            '5.1. За просрочку оплаты Поставщик вправе начислить неустойку в размере 1% от суммы задолженности за каждый день просрочки, но не более 10%.',
            '5.2. Стороны освобождаются от ответственности при наступлении обстоятельств непреодолимой силы.',
        ]),
        ('6. ФОРС-МАЖОР', [
            '6.1. Сторона, для которой возникли обстоятельства непреодолимой силы, обязана уведомить другую сторону в течение 3 (трёх) рабочих дней.',
            '6.2. Если форс-мажор продолжается более 6 (шести) месяцев подряд, каждая из сторон вправе расторгнуть договор в одностороннем порядке.',
        ]),
        ('7. ПОРЯДОК РАЗРЕШЕНИЯ СПОРОВ', [
            '7.1. Все споры стороны разрешают путём переговоров в течение 10 (десяти) рабочих дней.',
            '7.2. При недостижении согласия спор передаётся в арбитражный суд по месту нахождения ответчика.',
        ]),
        ('8. СРОК ДЕЙСТВИЯ ДОГОВОРА', [
            '8.1. Договор вступает в силу с даты подписания и действует до {{contract_end}}, с автоматической пролонгацией на следующий год при отсутствии уведомления о расторжении.',
        ]),
        ('9. ПРОЧИЕ УСЛОВИЯ', [
            '9.1. Договор составлен в 2 (двух) экземплярах, по одному для каждой из сторон.',
            '9.2. Все изменения к договору оформляются дополнительными соглашениями в письменной форме.',
            '9.3. По всему, что не урегулировано настоящим договором, стороны руководствуются действующим законодательством Российской Федерации.',
        ]),
    ]

    for title_text, paras in _sections:
        tp = doc.add_paragraph()
        tp.add_run(title_text).bold = True
        for pt in paras:
            doc.add_paragraph(pt)
        doc.add_paragraph()

    # Реквизиты
    sig = doc.add_paragraph()
    sig.add_run('РЕКВИЗИТЫ И ПОДПИСИ СТОРОН').bold = True

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    left  = tbl.cell(0, 0)
    right = tbl.cell(0, 1)

    left.text = (
        'ПОКУПАТЕЛЬ\n'
        '{{buyer_name}}\n'
        'ИНН: {{buyer_inn}}, КПП: {{buyer_kpp}}\n'
        'Адрес: {{buyer_address}}\n'
        'Р/С: {{buyer_rs}}\n'
        'Банк: {{buyer_bank}}\n'
        'БИК: {{buyer_bik}}, К/С: {{buyer_ks}}\n\n'
        '_________________ {{buyer_director}}\nМ.П.'
    )
    right.text = (
        f'ПОСТАВЩИК\n'
        f'{config.SELLER_NAME}\n'
        f'ИНН: {config.SELLER_INN}\n'
        f'Адрес: {config.SELLER_ADDRESS}\n'
        f'Р/С: {config.SELLER_RS}\n'
        f'Банк: {config.SELLER_BANK_NAME}\n'
        f'БИК: {config.SELLER_BIK}, К/С: {config.SELLER_KS}\n\n'
        f'_________________ {{{{seller_signature}}}}\nМ.П.'
    )

    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)


# ─────────────────────── EDO XML ──────────────────────────────────

def generate_xml(data: dict) -> bytes:
    """Генерирует XML-файл УПД в формате ФНС ON_NSCHFDOPPR v5.03 (windows-1251)."""
    items, total_excl, total_vat, total_incl = enrich_items(
        data['items'], Decimal(str(data.get('delivery', 0)))
    )
    doc_date = data['date']
    buyer    = data['buyer']
    num_raw  = data['doc_number']          # Б-000060

    file_id = (
        f"ON_NSCHFDOPPR_{buyer.get('inn','')}_"
        f"{buyer.get('kpp','')}_{config.SELLER_INN}_"
        f"{doc_date.strftime('%Y%m%d')}_{uuid.uuid4().hex[:8]}"
    )

    NSMAP = {
        'xs':  'http://www.w3.org/2001/XMLSchema',
        'xsi': 'http://www.w3.org/2001/XMLSchema-instance',
    }

    root = etree.Element('Файл', nsmap=NSMAP)
    root.set('ИдФайл',  file_id)
    root.set('ВерсФорм', '5.03')
    root.set('ВерсПрог', 'TelegramBot-DocFlow-1.4')

    doc_el = etree.SubElement(root, 'Документ')
    doc_el.set('КНД',            '1115131')
    doc_el.set('Функция',        'СЧФДОП')
    doc_el.set('НаимДокОпр',     'Счет-фактура и документ об отгрузке товаров (выполнении работ), '
                                  'передаче имущественных прав (документ об оказании услуг)')
    doc_el.set('ПоФактХЖ',       'Передаточный документ (акт)')
    doc_el.set('НомДок',         num_raw)
    doc_el.set('ДатаДок',        _date_short(doc_date))
    doc_el.set('НаимЭконСубСост', 'Индивидуальный предприниматель')
    doc_el.set('ВремФайл',       datetime.now().strftime('%H.%M.%S'))

    # ── СвСчФакт ──
    sv = etree.SubElement(doc_el, 'СвСчФакт')
    sv.set('НомерСчФ', num_raw)
    sv.set('ДатаСчФ',  _date_short(doc_date))

    # Продавец
    sv_prod = etree.SubElement(sv, 'СвПрод')
    id_sv   = etree.SubElement(sv_prod, 'ИдСв')
    ogrnip  = etree.SubElement(id_sv, 'ОГРНИП')
    ogrnip.set('ОГРНИП',       config.SELLER_OGRNIP)
    ogrnip.set('ДатаОГРНИП',  config.SELLER_OGRNIP_DATE)
    ogrnip.set('НаимОП',       config.SELLER_SHORT_NAME)
    fio = etree.SubElement(ogrnip, 'ФИО')
    fio.set('Фамилия',  'Шавкова')
    fio.set('Имя',      'Тамара')
    fio.set('Отчество', 'Расуловна')

    adr_prod = etree.SubElement(sv_prod, 'Адрес')
    adr_rf   = etree.SubElement(adr_prod, 'АдрРФ')
    adr_rf.set('КодРег',      config.SELLER_ADDR_REGION_CODE)
    adr_rf.set('НаимРегион',  config.SELLER_ADDR_REGION_NAME)

    rekv = etree.SubElement(sv_prod, 'РеквСчет')
    rekv.set('НомСч', config.SELLER_RS)
    bank_el = etree.SubElement(rekv, 'СвБанк')
    bank_el.set('НаимБанк', config.SELLER_BANK_NAME)
    bank_el.set('БИК',      config.SELLER_BIK)
    bank_el.set('КорСч',    config.SELLER_KS)

    # Грузоотправитель (он же)
    gruz_ot   = etree.SubElement(sv, 'ГрузОт')
    gruz_otpr = etree.SubElement(gruz_ot, 'ГрузОтпр')
    gruz_otpr.set('ОнЖе', 'он же')

    # Покупатель / Грузополучатель
    for tag in ('СвПокуп', 'ГрузПолуч'):
        sv_pok  = etree.SubElement(sv, tag)
        id_sv2  = etree.SubElement(sv_pok, 'ИдСв')
        sv_yul  = etree.SubElement(id_sv2, 'СвЮЛУч')
        sv_yul.set('НаимОрг', buyer.get('name', ''))
        sv_yul.set('ИННЮЛ',   buyer.get('inn', ''))
        if buyer.get('kpp'):
            sv_yul.set('КПП', buyer['kpp'])
        adr_pok = etree.SubElement(sv_pok, 'Адрес')
        adr_ino = etree.SubElement(adr_pok, 'АдрИно')
        adr_ino.set('КодСтр',   '643')
        adr_ino.set('НаимСтр',  'Россия')
        adr_ino.set('АдрТекст', buyer.get('address', ''))

    # Валюта
    dop = etree.SubElement(sv, 'ДопСвФХЖ1')
    dop.set('НаимОКВ', 'Российский рубль')
    dop.set('КодОКВ',  '643')
    dop.set('КурсВал', '1.00')

    # ── ТаблСчФакт ──
    tabl = etree.SubElement(doc_el, 'ТаблСчФакт')
    for idx, it in enumerate(items, 1):
        st = etree.SubElement(tabl, 'СведТов')
        st.set('НомСтр',      str(idx))
        st.set('НаимТов',     it['name'])
        if it.get('qty') not in ('', None):
            st.set('ОКЕИ_Тов', '796')
            st.set('НаимЕд',   str(it.get('unit', 'шт')))
            st.set('КолТов',   str(it['qty']))
            st.set('ЦенаТов',  f"{it['price_without_vat']:.2f}")
        st.set('СтТовБезНДС', f"{it['total_without_vat']:.2f}")
        st.set('НалСт',        config.VAT_LABEL)
        st.set('СтТовУчНал',  f"{it['total_with_vat']:.2f}")

        ak = etree.SubElement(st, 'АкцизТов')
        etree.SubElement(ak, 'БезАкциза').text = 'без акциза'

        nds = etree.SubElement(st, 'НДС')
        etree.SubElement(nds, 'СумНДС').text = f"{it['vat_amount']:.2f}"

    vsego = etree.SubElement(tabl, 'ВсегоОпл')
    vsego.set('СтТовБезНДСВсего', f"{total_excl:.2f}")
    vsego.set('СтТовУчНалВсего',  f"{total_incl:.2f}")
    vsego.set('КолЛистов', '1')
    nds_v = etree.SubElement(vsego, 'НДСВсего')
    etree.SubElement(nds_v, 'СумНДС').text = f"{total_vat:.2f}"

    # ── СвПодп ──
    sv_podp  = etree.SubElement(doc_el, 'СвПодп')
    podp_ip  = etree.SubElement(sv_podp, 'ПодпИП')
    podp_ip.set('ДатаПодп', _date_short(doc_date))
    fio2 = etree.SubElement(podp_ip, 'ФИО')
    fio2.set('Фамилия',  'Шавкова')
    fio2.set('Имя',      'Тамара')
    fio2.set('Отчество', 'Расуловна')
    sv_ip = etree.SubElement(podp_ip, 'СвИП')
    sv_ip.set('ОГРНИП',      config.SELLER_OGRNIP)
    sv_ip.set('ДатаОГРНИП', config.SELLER_OGRNIP_DATE)
    etree.SubElement(podp_ip, 'ОснПолн').text = 'ИП'

    return etree.tostring(
        root,
        encoding='windows-1251',
        xml_declaration=True,
        pretty_print=True,
    )
