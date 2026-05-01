"""
Telegram-бот «Автоматизация документооборота»
Работает в групповых чатах. Реагирует на /команды и @упоминания.
FSM-сессия привязана к пользователю и чату.
"""
import asyncio
import io
import json
import logging
import os
import re
from datetime import date
from decimal import Decimal

from aiogram import Bot, Dispatcher, F, Router
from aiogram.client.default import DefaultBotProperties
from aiogram.enums import ParseMode
from aiogram.filters import Command, CommandStart
from aiogram.types import BotCommand, BotCommandScopeDefault, BotCommandScopeAllGroupChats
from aiogram.fsm.context import FSMContext
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.types import (
    BufferedInputFile, Message,
    InlineKeyboardMarkup, InlineKeyboardButton,
    CallbackQuery, InputMediaDocument,
)

import config
import db
from parser import (InvoiceForm, Item, parse_invoice_text, validate_input, validate_parsed_invoice,
                    parse_buyer_card, parse_buyer_card_from_image, _looks_like_buyer_card)
from generator import (
    generate_invoice_pdf,
    generate_upd_pdf,
    generate_contract_pdf,
    generate_contract_docx,
    generate_xml,
)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(name)s: %(message)s',
)
log = logging.getLogger(__name__)

router = Router()
_bot_username: str = ''
_bot_id: int = 0


# ═══════════════════════════════════════════════════════════════
#  Helpers
# ═══════════════════════════════════════════════════════════════

def _items_from_parsed(parsed_items: list[Item]) -> list[dict]:
    return [{'name': i.name, 'qty': float(i.qty), 'unit': i.unit,
             'price': float(i.price)} for i in parsed_items]


def _fmt_items(items: list[dict]) -> str:
    lines = []
    for i, item in enumerate(items, 1):
        if item.get('qty'):
            lines.append(f"  {i}. {item['name']} × {item['qty']} {item['unit']} = {float(item['price']):.2f} руб.")
        else:
            lines.append(f"  {i}. {item['name']} = {float(item['price']):.2f} руб.")
    return '\n'.join(lines)


def _confirm_keyboard() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [
            InlineKeyboardButton(text="✏️ Реквизиты", callback_data="edit_buyer"),
            InlineKeyboardButton(text="✏️ Позиции",   callback_data="edit_items"),
        ],
        [
            InlineKeyboardButton(text="✅ Генерировать", callback_data="confirm_yes"),
            InlineKeyboardButton(text="❌ Отмена",       callback_data="confirm_no"),
        ],
    ])


def _confirm_text(data: dict) -> str:
    buyer    = data.get('buyer', {})
    items    = data.get('items', [])
    delivery = data.get('delivery', 0)
    lines    = [
        "📋 *Проверьте данные:*",
        f"\n*Покупатель:* {buyer.get('name', '—')} (ИНН {buyer.get('inn', '—')})",
        f"*Товары:*\n{_fmt_items(items)}",
    ]
    if delivery:
        lines.append(f"*Доставка:* {float(delivery):.2f} руб.")
    return '\n'.join(lines)


def _strip_mention(text: str) -> str:
    """Убирает @упоминание бота из начала строки."""
    if _bot_username:
        text = re.sub(rf'@{re.escape(_bot_username)}\s*', '', text, flags=re.IGNORECASE).strip()
    return text


_RE_BUYER_BLOCK_START = re.compile(
    r'^\s*(?:инн|кпп|бик|огрн|р\s*/\s*с|расчётный|расчетный'
    r'|кор\.?\s*счёт|корр\.?\s*счёт'
    r'|ип\s+\w|ооо\s+|зао\s+|пао\s+|ао\s+|оао\s+'
    r'|индивидуальный\s+предприниматель)',
    re.IGNORECASE,
)
_RE_CAPS_NAME = re.compile(r'^[А-ЯЁ\s«»"\']{10,}$')


def _split_items_buyer(text: str) -> tuple[str, str]:
    """Разделяет сообщение на часть с товарами и часть с реквизитами."""
    lines = text.splitlines()
    split_at = None
    for i, line in enumerate(lines):
        s = line.strip()
        if not s:
            continue
        if _RE_BUYER_BLOCK_START.match(s):
            # Шаг назад: если предыдущая непустая строка — заглавное имя контрагента
            j = i - 1
            while j >= 0 and not lines[j].strip():
                j -= 1
            if j >= 0 and _RE_CAPS_NAME.match(lines[j].strip()):
                split_at = j
            else:
                split_at = i
            break
    if split_at is not None:
        return '\n'.join(lines[:split_at]), '\n'.join(lines[split_at:])
    return text, ''


async def _try_parse_buyer_from_text(text: str, known_inn: str = '') -> dict | None:
    """Если в тексте есть карта реквизитов или хотя бы ИНН с кешем — возвращает buyer dict."""
    # Сначала проверяем кэш по известному ИНН
    if known_inn:
        cached = db.get_buyer(known_inn)
        if cached:
            return dict(cached)

    # Если текст похож на карту контрагента — парсим через Gemini
    if _looks_like_buyer_card(text):
        buyer = await parse_buyer_card(text, config.GEMINI_API_KEY)
        if buyer.get('inn'):
            cached = db.get_buyer(buyer['inn'])
            return dict(cached) if cached else buyer

    return None


_DEDICATED_COMMANDS = frozenset({
    'start', 'help', 'помощь',
    'cancel', 'отмена', 'stop',
    'invoice', 'счёт', 'счет', 'new',
    'edit_package',
})


def _is_trigger(text: str) -> bool:
    if not text:
        return False
    if text.startswith('/'):
        cmd = text[1:].split()[0].split('@')[0].lower()
        if cmd in _DEDICATED_COMMANDS:
            return False      # у этих команд есть свои хендлеры
        return True
    if _bot_username and f'@{_bot_username.lower()}' in text.lower():
        return True
    return False


# ═══════════════════════════════════════════════════════════════
#  Основной триггер (/команды и @упоминания)
# ═══════════════════════════════════════════════════════════════

@router.message(F.text.func(lambda t: _is_trigger(t or '')))
async def cmd_trigger(message: Message, state: FSMContext) -> None:
    raw_text  = message.text or ''
    clean     = _strip_mention(raw_text)          # убираем @mention
    clean     = re.sub(r'^/\S+\s*', '', clean).strip()  # убираем /команду

    # Если уже в FSM-сессии и это @упоминание (не /команда) —
    # передаём очищенный текст в текущий FSM-шаг вместо сброса
    current = await state.get_state()
    if current and not raw_text.startswith('/'):
        fake = message.model_copy(update={'text': clean or message.text})
        await _dispatch_fsm(fake, state, clean)
        return

    # Новая /команда — сбрасываем состояние и начинаем
    await state.clear()

    if not clean:
        await message.reply(
            "Введите описание заказа, например:\n"
            "`Подставка DeWalt DWX726 1 шт 30780 доставка 2160 ИНН 7735558789`",
            parse_mode=ParseMode.MARKDOWN,
        )
        await state.set_state(InvoiceForm.items)
        return

    items_text, buyer_text = _split_items_buyer(clean)
    parsed = await parse_invoice_text(items_text or clean, config.GEMINI_API_KEY)

    if parsed.items:
        await state.update_data(items=_items_from_parsed(parsed.items),
                                delivery=float(parsed.delivery))

        # Если в том же сообщении есть карта реквизитов — парсим сразу
        buyer = await _try_parse_buyer_from_text(buyer_text or clean, parsed.buyer_inn)
        if buyer:
            await state.update_data(buyer=buyer)
            await message.reply(_confirm_text(await state.get_data()),
                                parse_mode=ParseMode.MARKDOWN,
                                reply_markup=_confirm_keyboard())
            await state.set_state(InvoiceForm.confirm)
            return

        await message.reply(
            "ИНН покупателя?\n_Ответьте на это сообщение, если бот не видит ваши ответы_",
            parse_mode=ParseMode.MARKDOWN,
        )
        await state.set_state(InvoiceForm.buyer_inn)
    else:
        await message.reply(
            "Не распознал товары. Введите в формате:\n"
            "`Название 1 шт 30780 доставка 2160`",
            parse_mode=ParseMode.MARKDOWN,
        )
        await state.set_state(InvoiceForm.items)


async def _dispatch_fsm(message: Message, state: FSMContext, text: str) -> None:
    """Маршрутизатор FSM-шагов при @упоминании во время активной сессии."""
    current = await state.get_state()
    if current == InvoiceForm.items.state:
        await fsm_items(message, state)
    elif current == InvoiceForm.buyer_inn.state:
        await fsm_buyer_inn(message, state)
    elif current == InvoiceForm.buyer_name.state:
        await fsm_buyer_name(message, state)
    elif current == InvoiceForm.buyer_kpp.state:
        await fsm_buyer_kpp(message, state)
    elif current == InvoiceForm.buyer_address.state:
        await fsm_buyer_address(message, state)
    elif current == InvoiceForm.buyer_director.state:
        await fsm_buyer_director(message, state)
    elif current == InvoiceForm.buyer_bank.state:
        await fsm_buyer_bank(message, state)
    elif current == InvoiceForm.confirm.state:
        await fsm_confirm(message, state)
    elif current == InvoiceForm.edit_buyer.state:
        await fsm_edit_buyer(message, state)
    elif current == InvoiceForm.edit_items.state:
        await fsm_edit_items(message, state)


# ═══════════════════════════════════════════════════════════════
#  FSM — также принимаем ОТВЕТЫ (reply) на сообщения бота
#  (работает даже при включённом privacy mode)
# ═══════════════════════════════════════════════════════════════

def _is_reply_to_bot(message: Message) -> bool:
    return bool(
        message.reply_to_message and
        message.reply_to_message.from_user and
        message.reply_to_message.from_user.id == _bot_id
    )


@router.message(
    F.text,
    lambda msg: _is_reply_to_bot(msg),
)
async def handle_reply_to_bot(message: Message, state: FSMContext) -> None:
    """Обрабатывает ответы (reply) на сообщения бота — для работы с privacy mode."""
    await _dispatch_fsm(message, state, message.text or '')


# ═══════════════════════════════════════════════════════════════
#  FSM — сбор товаров
# ═══════════════════════════════════════════════════════════════

_not_cmd = ~F.text.startswith('/')


@router.message(InvoiceForm.items, _not_cmd)
async def fsm_items(message: Message, state: FSMContext) -> None:
    text             = _strip_mention(message.text or '')

    # Валидация входа перед Gemini (защита от DoS и prompt injection)
    is_valid, error_msg = validate_input(text,
                                         max_length=config.MAX_INPUT_LENGTH,
                                         max_lines=config.MAX_LINES)
    if not is_valid:
        await message.reply(error_msg)
        return

    items_text, buyer_text = _split_items_buyer(text)
    parsed = await parse_invoice_text(items_text or text, config.GEMINI_API_KEY)

    # Валидация выхода от Gemini (JSON schema validation)
    is_valid, error_msg = validate_parsed_invoice(parsed,
                                                   max_price=config.MAX_PRICE_PER_ITEM,
                                                   max_items=config.MAX_ITEMS_COUNT)
    if not is_valid:
        await message.reply(f"Ошибка парсинга: {error_msg}\n\n_Попробуйте переформулировать запрос._",
                           parse_mode=ParseMode.MARKDOWN)
        return

    if not parsed.items:
        await message.reply(
            "Не распознал. Пример:\n`Кабель ВВГ 50 м 15000 доставка 500`",
            parse_mode=ParseMode.MARKDOWN,
        )
        return
    await state.update_data(items=_items_from_parsed(parsed.items),
                            delivery=float(parsed.delivery))

    buyer = await _try_parse_buyer_from_text(buyer_text or text, parsed.buyer_inn)
    if buyer:
        await state.update_data(buyer=buyer)
        await message.reply(_confirm_text(await state.get_data()),
                            parse_mode=ParseMode.MARKDOWN,
                            reply_markup=_confirm_keyboard())
        await state.set_state(InvoiceForm.confirm)
        return

    await message.reply(
        "ИНН покупателя?\n_Отвечайте ОТВЕТОМ на мои сообщения_",
        parse_mode=ParseMode.MARKDOWN,
    )
    await state.set_state(InvoiceForm.buyer_inn)


# ═══════════════════════════════════════════════════════════════
#  FSM — ИНН покупателя
# ═══════════════════════════════════════════════════════════════

@router.message(InvoiceForm.buyer_inn, _not_cmd)
async def fsm_buyer_inn(message: Message, state: FSMContext) -> None:
    text = _strip_mention(message.text or '').strip()

    # Если вставили полную карту реквизитов — парсим через Gemini
    if _looks_like_buyer_card(text):
        wait = await message.reply("🔍 Читаю реквизиты…")
        buyer = await parse_buyer_card(text, config.GEMINI_API_KEY)
        if buyer.get('inn'):
            cached = db.get_buyer(buyer['inn'])
            if cached:
                buyer = dict(cached)
            await state.update_data(buyer=buyer)
            await wait.edit_text(
                f"✅ {buyer.get('name','?')} (ИНН {buyer['inn']})"
            )
            await message.reply(_confirm_text(await state.get_data()),
                                parse_mode=ParseMode.MARKDOWN,
                                reply_markup=_confirm_keyboard())
            await state.set_state(InvoiceForm.confirm)
        else:
            await wait.edit_text("❌ ИНН не найден. Введите ИНН отдельной строкой:")
        return

    inn_match = re.search(r'\d{10,12}', text)

    if not inn_match and text.lower() not in ('нет', 'no', '-'):
        await message.reply("Введите ИНН (10 или 12 цифр):")
        return

    inn = inn_match.group() if inn_match else ''
    if inn:
        buyer = db.get_buyer(inn)
        if buyer:
            await state.update_data(buyer=dict(buyer))
            await message.reply(_confirm_text(await state.get_data()),
                                parse_mode=ParseMode.MARKDOWN,
                                reply_markup=_confirm_keyboard())
            await state.set_state(InvoiceForm.confirm)
            return
        await state.update_data(buyer={'inn': inn})

    await message.reply("Полное наименование покупателя (ООО / ИП / ...):")
    await state.set_state(InvoiceForm.buyer_name)


# ═══════════════════════════════════════════════════════════════
#  FSM — данные покупателя
# ═══════════════════════════════════════════════════════════════

@router.message(InvoiceForm.buyer_name, _not_cmd)
async def fsm_buyer_name(message: Message, state: FSMContext) -> None:
    data  = await state.get_data()
    buyer = data.get('buyer', {})
    buyer['name'] = _strip_mention(message.text or '').strip()
    await state.update_data(buyer=buyer)
    await message.reply("КПП покупателя (или `-` если нет):")
    await state.set_state(InvoiceForm.buyer_kpp)


@router.message(InvoiceForm.buyer_kpp, _not_cmd)
async def fsm_buyer_kpp(message: Message, state: FSMContext) -> None:
    data  = await state.get_data()
    buyer = data.get('buyer', {})
    t = _strip_mention(message.text or '').strip()
    buyer['kpp'] = '' if t == '-' else t
    await state.update_data(buyer=buyer)
    await message.reply("Юридический адрес покупателя:")
    await state.set_state(InvoiceForm.buyer_address)


@router.message(InvoiceForm.buyer_address, _not_cmd)
async def fsm_buyer_address(message: Message, state: FSMContext) -> None:
    data  = await state.get_data()
    buyer = data.get('buyer', {})
    buyer['address'] = _strip_mention(message.text or '').strip()
    await state.update_data(buyer=buyer)
    await message.reply("ФИО руководителя (или `-`):")
    await state.set_state(InvoiceForm.buyer_director)


@router.message(InvoiceForm.buyer_director, _not_cmd)
async def fsm_buyer_director(message: Message, state: FSMContext) -> None:
    data  = await state.get_data()
    buyer = data.get('buyer', {})
    t = _strip_mention(message.text or '').strip()
    buyer['director'] = '' if t == '-' else t
    await state.update_data(buyer=buyer)
    await message.reply(
        "Банковские реквизиты (каждый с новой строки):\n"
        "```\nР/С: 40702810xxxxxxxxx\nБанк: ПАО Сбербанк\nБИК: 044525225\nК/С: 30101810400000000225\n```",
        parse_mode=ParseMode.MARKDOWN,
    )
    await state.set_state(InvoiceForm.buyer_bank)


@router.message(InvoiceForm.buyer_bank, _not_cmd)
async def fsm_buyer_bank(message: Message, state: FSMContext) -> None:
    data  = await state.get_data()
    buyer = data.get('buyer', {})
    text  = _strip_mention(message.text or '')

    def _find(pattern: str) -> str:
        m = re.search(pattern, text, re.IGNORECASE)
        return m.group(1).strip() if m else ''

    buyer['rs'] = _find(
        r'(?:[рp][/\/][сc]|расч[её]тный\s+сч[её]т|расчетный\s+счет)'
        r'\s*[:\-]?\s*([\d]+)'
    )
    buyer['bank_name'] = _find(r'банк\s*[:\-]?\s*(.+)')
    buyer['bik']       = _find(r'бик\s*[:\-]?\s*(\d+)')
    buyer['ks'] = _find(
        r'(?:к[/\/]с|корр?\.?\s*сч[её]т|корреспондентский\s+сч[её]т)'
        r'\s*[:\-]?\s*([\d]+)'
    )

    await state.update_data(buyer=buyer)
    await message.reply(_confirm_text(await state.get_data()),
                        parse_mode=ParseMode.MARKDOWN,
                        reply_markup=_confirm_keyboard())
    await state.set_state(InvoiceForm.confirm)


# ═══════════════════════════════════════════════════════════════
#  FSM — подтверждение и генерация
# ═══════════════════════════════════════════════════════════════

_YES = {'да', 'yes', 'ок', 'ok', 'подтвердить', '✅', '+', 'generate', 'да!'}


@router.callback_query(F.data == "confirm_yes")
async def cb_confirm_yes(callback: CallbackQuery, state: FSMContext) -> None:
    await callback.answer()
    try:
        await callback.message.edit_reply_markup(reply_markup=None)
    except Exception:
        pass
    await _send_docs(callback.message, state, from_user_id=callback.from_user.id)


@router.callback_query(F.data == "confirm_no")
async def cb_confirm_no(callback: CallbackQuery, state: FSMContext) -> None:
    await callback.answer("Отменено")
    await state.clear()
    try:
        await callback.message.edit_text("❌ Отменено.")
    except Exception:
        await callback.message.reply("❌ Отменено.")


@router.message(InvoiceForm.confirm, _not_cmd)
async def fsm_confirm(message: Message, state: FSMContext) -> None:
    text = _strip_mention(message.text or '').strip().lower().lstrip('/')
    if text in _YES:
        await _send_docs(message, state)
    elif text in {'нет', 'no', 'отмена', 'cancel', '-'}:
        await state.clear()
        await message.reply("❌ Отменено.")
    else:
        await message.reply(
            "Используйте кнопки выше или ответьте `да` / `отмена`.",
            parse_mode=ParseMode.MARKDOWN,
        )


# ═══════════════════════════════════════════════════════════════
#  Редактирование реквизитов (кнопка ✏️ Реквизиты)
# ═══════════════════════════════════════════════════════════════

@router.callback_query(F.data.startswith("pkg_"))
async def cb_load_package(callback: CallbackQuery, state: FSMContext) -> None:
    """Загружает пакет по нажатию кнопки в /edit_package."""
    await callback.answer()
    try:
        pkg_id = int(callback.data.split("_", 1)[1])
    except (ValueError, IndexError):
        await callback.message.reply("❌ Неверный ID пакета.")
        return

    pkg = db.get_package(pkg_id)
    if not pkg:
        await callback.message.reply(f"❌ Пакет `#{pkg_id}` не найден.", parse_mode=ParseMode.MARKDOWN)
        return
    try:
        data = json.loads(pkg['json_data'])
    except Exception:
        await callback.message.reply("❌ Ошибка чтения данных пакета.")
        return

    await state.clear()
    await state.update_data(**data)
    buyer_name = data.get('buyer', {}).get('name', '') or data.get('buyer', {}).get('inn', '?')
    try:
        await callback.message.edit_reply_markup(reply_markup=None)
    except Exception:
        pass
    await callback.message.reply(
        f"📦 Пакет `#{pkg_id}` загружен ({buyer_name})\n"
        f"Проверьте данные и нажмите *Генерировать* или отредактируйте:",
        parse_mode=ParseMode.MARKDOWN,
    )
    await callback.message.reply(
        _confirm_text(data),
        parse_mode=ParseMode.MARKDOWN,
        reply_markup=_confirm_keyboard(),
    )
    await state.set_state(InvoiceForm.confirm)


@router.callback_query(F.data == "edit_buyer")
async def cb_edit_buyer(callback: CallbackQuery, state: FSMContext) -> None:
    await callback.answer()
    try:
        await callback.message.edit_reply_markup(reply_markup=None)
    except Exception:
        pass
    await callback.message.reply(
        "✏️ *Редактирование реквизитов*\n\n"
        "Отправьте один из вариантов:\n"
        "• Карту контрагента текстом (ИНН, КПП, банк…)\n"
        "• Фото документа с реквизитами\n"
        "• Файл .docx / .txt / .pdf\n"
        "• Только ИНН — если есть в базе",
        parse_mode=ParseMode.MARKDOWN,
    )
    await state.set_state(InvoiceForm.edit_buyer)


@router.message(InvoiceForm.edit_buyer, _not_cmd)
async def fsm_edit_buyer(message: Message, state: FSMContext) -> None:
    text = _strip_mention(message.text or '').strip()
    wait = await message.reply("🔍 Обновляю реквизиты…")

    # Попытка разобрать как карту реквизитов
    if _looks_like_buyer_card(text):
        buyer = await parse_buyer_card(text, config.GEMINI_API_KEY)
        if buyer.get('inn'):
            cached = db.get_buyer(buyer['inn'])
            if cached:
                buyer = dict(cached)
            await state.update_data(buyer=buyer)
            await wait.edit_text(f"✅ {buyer.get('name', '?')} (ИНН {buyer['inn']})")
            await message.reply(_confirm_text(await state.get_data()),
                                parse_mode=ParseMode.MARKDOWN,
                                reply_markup=_confirm_keyboard())
            await state.set_state(InvoiceForm.confirm)
            return

    # Попытка найти ИНН в тексте и загрузить из базы
    inn_match = re.search(r'\d{10,12}', text)
    if inn_match:
        inn = inn_match.group()
        buyer = db.get_buyer(inn)
        if buyer:
            await state.update_data(buyer=dict(buyer))
            await wait.edit_text(f"✅ {buyer.get('name', '?')} (ИНН {inn})")
            await message.reply(_confirm_text(await state.get_data()),
                                parse_mode=ParseMode.MARKDOWN,
                                reply_markup=_confirm_keyboard())
            await state.set_state(InvoiceForm.confirm)
            return

    await wait.edit_text(
        "❌ Не удалось распознать реквизиты.\n"
        "Отправьте карту контрагента, фото, файл или ИНН."
    )


# ═══════════════════════════════════════════════════════════════
#  Редактирование позиций (кнопка ✏️ Позиции)
# ═══════════════════════════════════════════════════════════════

@router.callback_query(F.data == "edit_items")
async def cb_edit_items(callback: CallbackQuery, state: FSMContext) -> None:
    await callback.answer()
    try:
        await callback.message.edit_reply_markup(reply_markup=None)
    except Exception:
        pass
    data  = await state.get_data()
    items = data.get('items', [])

    lines = ["✏️ *Редактирование позиций:*\n"]
    for i, item in enumerate(items, 1):
        lines.append(
            f"{i}. {item['name']}\n"
            f"   × {item['qty']} {item['unit']} × {float(item['price']):.2f} руб."
        )

    lines.append(
        "\nВведите: `НОМЕР количество цена`\n"
        "Пример: `1 2 11880` — позиция 1, кол-во 2, цена 11 880 руб.\n"
        "Только цена: `1 11880`"
    )
    await callback.message.reply('\n'.join(lines), parse_mode=ParseMode.MARKDOWN)
    await state.set_state(InvoiceForm.edit_items)


@router.message(InvoiceForm.edit_items, _not_cmd)
async def fsm_edit_items(message: Message, state: FSMContext) -> None:
    text       = _strip_mention(message.text or '').strip()
    text_lower = text.lower()
    data       = await state.get_data()
    items      = list(data.get('items', []))

    # ── Удалить последнюю позицию ──────────────────────────────
    if re.search(r'удал[иь]\s+последн|del\s+last|remove\s+last', text_lower):
        if items:
            removed = items.pop()
            await state.update_data(items=items)
            await message.reply(
                f"🗑 Удалена: *{removed['name']}*",
                parse_mode=ParseMode.MARKDOWN,
            )
            await message.reply(
                _confirm_text(await state.get_data()),
                parse_mode=ParseMode.MARKDOWN,
                reply_markup=_confirm_keyboard(),
            )
            await state.set_state(InvoiceForm.confirm)
        else:
            await message.reply("Список позиций уже пуст.")
        return

    # ── Удалить по номеру: «удали 9», «del 9», «-9» ───────────
    del_m = re.search(r'(?:удал[иь]|удалить|del(?:ete)?|remove|-)\s*(\d+)', text_lower)
    if del_m:
        num = int(del_m.group(1))
        if 1 <= num <= len(items):
            removed = items.pop(num - 1)
            await state.update_data(items=items)
            await message.reply(
                f"🗑 Удалена позиция {num}: *{removed['name']}*",
                parse_mode=ParseMode.MARKDOWN,
            )
            await message.reply(
                _confirm_text(await state.get_data()),
                parse_mode=ParseMode.MARKDOWN,
                reply_markup=_confirm_keyboard(),
            )
            await state.set_state(InvoiceForm.confirm)
        else:
            await message.reply(f"Нет позиции #{num}. Доступно: 1–{len(items)}.")
        return

    # ── Числовой формат: N [qty] price ────────────────────────
    parts = text.split()
    if not parts:
        await message.reply(
            "Введите: `НОМЕР количество цена`\nПример: `1 2 11880`",
            parse_mode=ParseMode.MARKDOWN,
        )
        return

    try:
        num = int(parts[0].rstrip('.'))
    except ValueError:
        await message.reply(
            "Не понял команду. Варианты:\n"
            "• `1 2 11880` — позиция 1, кол-во 2, цена 11 880 руб.\n"
            "• `1 11880` — изменить только цену\n"
            "• `удали 9` — удалить позицию 9\n"
            "• `удали последнюю` — удалить последнюю позицию",
            parse_mode=ParseMode.MARKDOWN,
        )
        return

    if not 1 <= num <= len(items):
        await message.reply(f"Номер позиции от 1 до {len(items)}.")
        return

    try:
        if len(parts) >= 3:
            items[num - 1]['qty']   = float(parts[1].replace(',', '.'))
            items[num - 1]['price'] = float(parts[2].replace(',', '.'))
        elif len(parts) == 2:
            items[num - 1]['price'] = float(parts[1].replace(',', '.'))
        else:
            await message.reply(
                "Укажите хотя бы цену: `1 11880`",
                parse_mode=ParseMode.MARKDOWN,
            )
            return

        await state.update_data(items=items)
        await message.reply(
            _confirm_text(await state.get_data()),
            parse_mode=ParseMode.MARKDOWN,
            reply_markup=_confirm_keyboard(),
        )
        await state.set_state(InvoiceForm.confirm)

    except (ValueError, IndexError):
        await message.reply(
            "Неверный формат числа. Пример: `1 2 11880`",
            parse_mode=ParseMode.MARKDOWN,
        )


# ═══════════════════════════════════════════════════════════════
#  Генерация документов
# ═══════════════════════════════════════════════════════════════

async def _send_docs(message: Message, state: FSMContext, from_user_id: int | None = None) -> None:
    data = await state.get_data()
    await state.clear()

    serial, doc_number = db.next_doc_number()
    buyer    = data['buyer']
    items    = data['items']
    delivery = Decimal(str(data.get('delivery', 0)))

    gen_data = {
        'doc_number': doc_number,
        'date':       date.today(),
        'buyer':      buyer,
        'items':      items,
        'delivery':   delivery,
    }

    status = await message.answer(f"⚙️ Генерирую пакет *{doc_number}*…",
                                  parse_mode=ParseMode.MARKDOWN)

    errors: list[str] = []
    media:  list[InputMediaDocument] = []

    # ── Счёт PDF ──
    try:
        b = generate_invoice_pdf(gen_data)
        media.append(InputMediaDocument(
            media=BufferedInputFile(b, f"Счёт_{doc_number}.pdf"),
            caption=(f"📦 *Пакет {doc_number}*\n"
                     f"📄 Счёт на оплату\n"
                     f"📄 УПД Статус 1\n"
                     f"📝 Договор поставки\n"
                     f"🗂 XML для ЭДО"),
            parse_mode=ParseMode.MARKDOWN,
        ))
    except Exception as e:
        log.exception("Invoice PDF error")
        errors.append(f"Счёт: {e}")

    # ── УПД PDF ──
    try:
        b = generate_upd_pdf(gen_data)
        media.append(InputMediaDocument(media=BufferedInputFile(b, f"УПД_{doc_number}.pdf")))
    except Exception as e:
        log.exception("UPD PDF error")
        errors.append(f"УПД: {e}")

    # ── Договор PDF (DOCX как резерв) ──
    try:
        b = generate_contract_pdf(gen_data)
        media.append(InputMediaDocument(media=BufferedInputFile(b, f"Договор_{doc_number}.pdf")))
    except Exception as e:
        log.exception("Contract PDF error, trying DOCX fallback")
        try:
            b = generate_contract_docx(gen_data)
            media.append(InputMediaDocument(media=BufferedInputFile(b, f"Договор_{doc_number}.docx")))
        except Exception as e2:
            log.exception("Contract DOCX error")
            errors.append(f"Договор: {e2}")

    # ── XML ──
    try:
        b = generate_xml(gen_data)
        fname = (f"ON_NSCHFDOPPR_{buyer.get('inn','')}_{buyer.get('kpp','')}_"
                 f"{config.SELLER_INN}_{doc_number}.xml")
        media.append(InputMediaDocument(media=BufferedInputFile(b, fname)))
    except Exception as e:
        log.exception("XML error")
        errors.append(f"XML: {e}")

    # ── Отправка одним постом ──
    try:
        await status.delete()
    except Exception:
        pass

    if media:
        await message.answer_media_group(media)

    if errors:
        await message.answer("⚠️ Ошибки при генерации:\n" + '\n'.join(errors))

    # ── Сохранение в БД ──
    try:
        user_id = from_user_id or (message.from_user.id if message.from_user else 0)
        total   = sum(float(i['price']) * float(i.get('qty') or 1) for i in items) + float(delivery)
        db.save_invoice(doc_number, serial, buyer.get('inn', ''), items,
                        float(delivery), total, message.chat.id, user_id)
        if buyer.get('inn'):
            db.upsert_buyer(buyer)

        # Сохраняем пакет для возможности повторного использования
        pkg_data = {'items': items, 'buyer': buyer, 'delivery': float(delivery)}
        pkg_id   = db.save_package(user_id, buyer.get('inn', ''), pkg_data)
        await message.answer(
            f"💾 Пакет сохранён `#{pkg_id}` — для повторной генерации: `/edit_package {pkg_id}`",
            parse_mode=ParseMode.MARKDOWN,
        )
    except Exception as e:
        log.exception("DB save error: %s", e)


# ═══════════════════════════════════════════════════════════════
#  Загрузка файла с реквизитами (.docx / .txt / .pdf)
# ═══════════════════════════════════════════════════════════════

async def _extract_file_text(bot: Bot, document) -> str:
    """Скачивает файл из Telegram и возвращает plain-text."""
    file = await bot.get_file(document.file_id)
    buf = io.BytesIO()
    await bot.download_file(file.file_path, buf)
    buf.seek(0)
    name = (document.file_name or '').lower()

    if name.endswith('.txt'):
        raw = buf.read()
        for enc in ('utf-8', 'cp1251', 'latin-1'):
            try:
                return raw.decode(enc)
            except UnicodeDecodeError:
                pass
        return raw.decode('utf-8', errors='replace')

    if name.endswith('.docx'):
        from docx import Document as DocxDoc
        doc = DocxDoc(buf)
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                line = ' | '.join(cells)
                if line.strip(' |'):
                    lines.append(line)
        return '\n'.join(lines)

    if name.endswith('.pdf'):
        try:
            import pypdf
            reader = pypdf.PdfReader(buf)
            return '\n'.join(page.extract_text() or '' for page in reader.pages)
        except ImportError:
            pass
        try:
            from pdfminer.high_level import extract_text as pdfminer_extract
            buf.seek(0)
            return pdfminer_extract(buf)
        except ImportError:
            pass

    return ''


@router.message(F.photo)
async def handle_photo(message: Message, state: FSMContext) -> None:
    """Принимает фото/скриншот с реквизитами покупателя — читает через Gemini Vision."""
    if not config.GEMINI_API_KEY:
        await message.reply("❌ Gemini API не настроен — распознавание фото недоступно.")
        return

    wait_msg = await message.reply("🔍 Читаю реквизиты с фото…")

    try:
        photo = message.photo[-1]
        file  = await message.bot.get_file(photo.file_id)
        buf   = io.BytesIO()
        await message.bot.download_file(file.file_path, buf)
        image_bytes = buf.getvalue()
    except Exception as e:
        await wait_msg.edit_text(f"❌ Не удалось скачать фото: {e}")
        return

    buyer = await parse_buyer_card_from_image(image_bytes, config.GEMINI_API_KEY)

    if not buyer.get('inn'):
        await wait_msg.edit_text(
            "❌ ИНН не найден на фото. Убедитесь, что на снимке видны реквизиты покупателя."
        )
        return

    cached = db.get_buyer(buyer['inn'])
    if cached:
        buyer = dict(cached)

    await state.update_data(buyer=buyer)
    data = await state.get_data()

    b = buyer
    preview = (
        f"✅ *Реквизиты с фото:*\n"
        f"*{b.get('name') or '—'}*\n"
        f"ИНН: `{b.get('inn','—')}` / КПП: `{b.get('kpp','—')}`\n"
        f"Адрес: {b.get('address') or '—'}\n"
        f"Р/С: `{b.get('rs','—')}`\n"
        f"Банк: {b.get('bank_name') or '—'} (БИК `{b.get('bik','—')}`)"
    )
    await wait_msg.edit_text(preview, parse_mode=ParseMode.MARKDOWN)

    if data.get('items'):
        await message.reply(_confirm_text(data), parse_mode=ParseMode.MARKDOWN,
                            reply_markup=_confirm_keyboard())
        await state.set_state(InvoiceForm.confirm)
    else:
        await message.reply(
            "Реквизиты сохранены. Теперь введите товары:\n"
            "`Название 1 шт 30780 доставка 2160`",
            parse_mode=ParseMode.MARKDOWN,
        )
        await state.set_state(InvoiceForm.items)


@router.message(F.document)
async def handle_document(message: Message, state: FSMContext) -> None:
    """Принимает файл с реквизитами покупателя (.docx / .txt / .pdf)."""
    doc = message.document
    if not doc:
        return
    name = (doc.file_name or '').lower()
    if not any(name.endswith(ext) for ext in ('.docx', '.txt', '.pdf')):
        return

    wait_msg = await message.reply("📂 Читаю реквизиты из файла…")

    try:
        text = await _extract_file_text(message.bot, doc)
    except Exception as e:
        log.exception("File extract error")
        await wait_msg.edit_text(f"❌ Ошибка чтения файла: {e}")
        return

    if not text.strip():
        await wait_msg.edit_text("❌ Не удалось извлечь текст. Пришлите .txt или .docx файл.")
        return

    buyer = await parse_buyer_card(text, config.GEMINI_API_KEY)
    if not buyer.get('inn'):
        await wait_msg.edit_text(
            "❌ ИНН не найден в файле. Убедитесь, что файл содержит реквизиты покупателя."
        )
        return

    cached = db.get_buyer(buyer['inn'])
    if cached:
        buyer = dict(cached)

    await state.update_data(buyer=buyer)
    data = await state.get_data()

    b = buyer
    preview = (
        f"✅ *Реквизиты прочитаны:*\n"
        f"*{b.get('name') or '—'}*\n"
        f"ИНН: `{b.get('inn','—')}` / КПП: `{b.get('kpp','—')}`\n"
        f"Адрес: {b.get('address') or '—'}\n"
        f"Р/С: `{b.get('rs','—')}`\n"
        f"Банк: {b.get('bank_name') or '—'} (БИК `{b.get('bik','—')}`)"
    )
    await wait_msg.edit_text(preview, parse_mode=ParseMode.MARKDOWN)

    if data.get('items'):
        await message.reply(_confirm_text(data), parse_mode=ParseMode.MARKDOWN, reply_markup=_confirm_keyboard())
        await state.set_state(InvoiceForm.confirm)
    else:
        await message.reply(
            "Реквизиты сохранены. Теперь введите товары:\n"
            "`Название 1 шт 30780 доставка 2160`",
            parse_mode=ParseMode.MARKDOWN,
        )
        await state.set_state(InvoiceForm.items)


# ═══════════════════════════════════════════════════════════════
#  Служебные команды
# ═══════════════════════════════════════════════════════════════

_HELP_TEXT = (
    "📄 *FineBuh Doc Bot v2 — автоматический документооборот*\n"
    "Продавец: ИП Шавкова Тамара Расуловна · НДС 5% (ОСНО)\n"
    "🤖 Разбор позиций и реквизитов — _Google Gemini AI_\n\n"

    "━━━━━━━━━━━━━━━━━━━━━\n"
    "📦 *Что получаете на каждый счёт*\n"
    "━━━━━━━━━━━━━━━━━━━━━\n"
    "  📄 Счёт на оплату — PDF\n"
    "  📄 УПД Статус 1 — PDF\n"
    "  📝 Договор поставки — PDF\n"
    "  🗂 XML для ЭДО (формат ФНС)\n"
    "Все 4 файла приходят одним сообщением.\n\n"

    "━━━━━━━━━━━━━━━━━━━━━\n"
    "🚀 *Как выставить счёт*\n"
    "━━━━━━━━━━━━━━━━━━━━━\n"
    "1\\. Напишите `/invoice` — бот попросит список товаров\n"
    "2\\. Отправьте позиции в любом формате:\n\n"
    "`Кабель ВВГ 50 м 15000 доставка 500`\n"
    "или нумерованным списком:\n"
    "`1. Шуруповерт DeWalt DCD708 - 1шт - 21600 руб`\n"
    "`2. Набор бит 37шт flextorq - 1шт - 2808 руб`\n\n"
    "AI разберёт любой формат — даже если в названии есть числа или единицы.\n"
    "Цены указывайте *с НДС 5%*.\n\n"

    "━━━━━━━━━━━━━━━━━━━━━\n"
    "🏢 *Реквизиты покупателя — 4 способа*\n"
    "━━━━━━━━━━━━━━━━━━━━━\n"
    "1️⃣ *ИНН* — бот найдёт реквизиты в кэше _(после первой генерации)_\n\n"
    "2️⃣ *Текст* — вставьте карту контрагента прямо в сообщение с товарами\n\n"
    "3️⃣ *Фото* — сфотографируйте карту / скриншот из 1С — Gemini Vision прочитает\n\n"
    "4️⃣ *Файл* — `.docx` `.txt` `.pdf` с реквизитами\n\n"

    "━━━━━━━━━━━━━━━━━━━━━\n"
    "✏️ *Редактирование перед генерацией*\n"
    "━━━━━━━━━━━━━━━━━━━━━\n"
    "На экране подтверждения — две кнопки:\n\n"
    "• *✏️ Реквизиты* — обновить данные покупателя (текст / фото / файл / ИНН)\n"
    "• *✏️ Позиции* — изменить количество, цену или удалить позицию\n\n"
    "В режиме редактирования позиций:\n"
    "  `1 2 11880` — позиция 1, кол-во 2, цена 11 880\n"
    "  `1 11880` — только цена\n"
    "  `удали 9` — удалить позицию №9\n"
    "  `удали последнюю` — удалить последнюю\n\n"

    "━━━━━━━━━━━━━━━━━━━━━\n"
    "💾 *История пакетов*\n"
    "━━━━━━━━━━━━━━━━━━━━━\n"
    "После каждой генерации пакет сохраняется автоматически:\n"
    "`/edit_package` — список последних пакетов\n"
    "`/edit_package 42` — загрузить пакет №42 для повторной / изменённой генерации\n\n"

    "━━━━━━━━━━━━━━━━━━━━━\n"
    "💬 *Команды*\n"
    "━━━━━━━━━━━━━━━━━━━━━\n"
    "`/invoice` — выставить счёт\n"
    "`/edit_package` — история пакетов\n"
    "`/cancel` — отменить текущую операцию\n"
    "`/help` — эта справка\n\n"

    "━━━━━━━━━━━━━━━━━━━━━\n"
    "💡 *Советы*\n"
    "━━━━━━━━━━━━━━━━━━━━━\n"
    "• В групповом чате — упоминайте `@FineBuh_Doc_Bot` или отвечайте на его сообщения\n"
    "• Повторный счёт тому же покупателю — достаточно ввести ИНН\n"
    "• Нумерация сквозная: Б-000001, Б-000002 …"
)


@router.message(CommandStart())
async def cmd_start(message: Message) -> None:
    await message.reply(_HELP_TEXT, parse_mode=ParseMode.MARKDOWN)


@router.message(Command('help', 'помощь'))
async def cmd_help(message: Message) -> None:
    await message.reply(_HELP_TEXT, parse_mode=ParseMode.MARKDOWN)


@router.message(Command('invoice', 'счёт', 'счет', 'new'))
async def cmd_invoice(message: Message, state: FSMContext) -> None:
    """Явная команда /invoice — псевдоним основного триггера."""
    raw   = message.text or ''
    clean = re.sub(r'^/\S+\s*', '', raw).strip()
    clean = _strip_mention(clean)
    await state.clear()

    if not clean:
        await message.reply(
            "Введите товары:\n"
            "`/invoice Название 1 шт 11880 доставка 500`\n\n"
            "Или отправьте описание следующим сообщением.",
            parse_mode=ParseMode.MARKDOWN,
        )
        await state.set_state(InvoiceForm.items)
        return

    items_text, buyer_text = _split_items_buyer(clean)
    parsed = await parse_invoice_text(items_text or clean, config.GEMINI_API_KEY)

    if parsed.items:
        await state.update_data(items=_items_from_parsed(parsed.items),
                                delivery=float(parsed.delivery))
        buyer = await _try_parse_buyer_from_text(buyer_text or clean, parsed.buyer_inn)
        if buyer:
            await state.update_data(buyer=buyer)
            await message.reply(_confirm_text(await state.get_data()),
                                parse_mode=ParseMode.MARKDOWN,
                                reply_markup=_confirm_keyboard())
            await state.set_state(InvoiceForm.confirm)
            return
        await message.reply(
            "ИНН покупателя?\n_Или отправьте фото / файл с реквизитами_",
            parse_mode=ParseMode.MARKDOWN,
        )
        await state.set_state(InvoiceForm.buyer_inn)
    else:
        await message.reply(
            "Не распознал товары. Пример:\n`/invoice Кабель ВВГ 50 м 15000`",
            parse_mode=ParseMode.MARKDOWN,
        )
        await state.set_state(InvoiceForm.items)


@router.message(Command('отмена', 'cancel', 'stop'))
async def cmd_cancel(message: Message, state: FSMContext) -> None:
    await state.clear()
    await message.reply("❌ Отменено.")


@router.message(Command('edit_package'))
async def cmd_edit_package(message: Message, state: FSMContext) -> None:
    """Загружает сохранённый пакет в FSM для повторной/изменённой генерации."""
    args = (message.text or '').split()
    user_id = message.from_user.id if message.from_user else 0

    if len(args) < 2:
        # Показываем список последних пакетов пользователя
        packages = db.list_packages(user_id)
        if not packages:
            await message.reply(
                "📭 Нет сохранённых пакетов.\n"
                "После каждой генерации пакет сохраняется автоматически.",
            )
            return
        buttons = []
        for p in packages:
            inn = p.get('client_inn') or '—'
            dt  = (p.get('created_at') or '')[:10]
            try:
                d     = json.loads(p['json_data'])
                buyer = d.get('buyer', {})
                inn   = buyer.get('inn') or inn
                name  = buyer.get('name', '') or ''
                name  = re.sub(
                    r'индивидуальный\s+предприниматель\s*',
                    'ИП ', name, flags=re.IGNORECASE
                ).strip()
                label = f"#{p['id']} {dt} · {name[:18]}" if name else f"#{p['id']} {dt} · ИНН {inn}"
            except Exception:
                label = f"#{p['id']} {dt} · ИНН {inn}"
            buttons.append([InlineKeyboardButton(
                text=label,
                callback_data=f"pkg_{p['id']}"
            )])
        kb = InlineKeyboardMarkup(inline_keyboard=buttons)
        await message.reply("📦 *Последние пакеты — нажмите чтобы загрузить:*",
                            parse_mode=ParseMode.MARKDOWN,
                            reply_markup=kb)
        return

    try:
        pkg_id = int(args[1])
    except ValueError:
        await message.reply("Укажите числовой ID. Пример: `/edit_package 42`",
                            parse_mode=ParseMode.MARKDOWN)
        return

    pkg = db.get_package(pkg_id)
    if not pkg:
        await message.reply(f"❌ Пакет `#{pkg_id}` не найден.", parse_mode=ParseMode.MARKDOWN)
        return

    try:
        data = json.loads(pkg['json_data'])
    except Exception:
        await message.reply("❌ Ошибка чтения данных пакета.")
        return

    await state.clear()
    await state.update_data(**data)
    buyer_name = data.get('buyer', {}).get('name', '') or data.get('buyer', {}).get('inn', '?')
    await message.reply(
        f"📦 Пакет `#{pkg_id}` загружен ({buyer_name})\n"
        f"Проверьте данные и нажмите *Генерировать* или отредактируйте:",
        parse_mode=ParseMode.MARKDOWN,
    )
    await message.reply(
        _confirm_text(data),
        parse_mode=ParseMode.MARKDOWN,
        reply_markup=_confirm_keyboard(),
    )
    await state.set_state(InvoiceForm.confirm)


# ═══════════════════════════════════════════════════════════════
#  Точка входа
# ═══════════════════════════════════════════════════════════════

async def main() -> None:
    global _bot_username, _bot_id

    # Загружаем .env
    try:
        from dotenv import load_dotenv
        load_dotenv(os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '.env'))
    except ImportError:
        pass

    if not config.BOT_TOKEN:
        raise RuntimeError("BOT_TOKEN не задан!")

    db.init_db()

    # Cleanup expired data on startup
    deleted = db.cleanup_expired()
    if sum(deleted) > 0:
        log.info(f"Cleaned up: {deleted[0]} buyers, {deleted[1]} invoices, {deleted[2]} packages")

    os.makedirs(config.OUTPUT_DIR, exist_ok=True)

    bot = Bot(
        token=config.BOT_TOKEN,
        default=DefaultBotProperties(parse_mode=ParseMode.HTML),
    )
    me = await bot.get_me()
    _bot_username = me.username or ''
    _bot_id       = me.id
    log.info("Бот запущен: @%s (id=%d)", _bot_username, _bot_id)

    # Регистрируем команды — появляются в меню "/" у пользователей
    commands = [
        BotCommand(command='invoice',      description='Выставить счёт'),
        BotCommand(command='edit_package', description='Список / загрузка пакетов'),
        BotCommand(command='cancel',       description='Отменить текущую операцию'),
        BotCommand(command='help',         description='Справка по боту'),
        BotCommand(command='start',        description='Начать работу'),
    ]
    await bot.set_my_commands(commands, scope=BotCommandScopeDefault())
    await bot.set_my_commands(commands, scope=BotCommandScopeAllGroupChats())

    # Описание бота (видно в профиле)
    try:
        await bot.set_my_description(
            "Автоматический документооборот на базе Google Gemini AI.\n"
            "Отправьте список товаров → получите Счёт + УПД + Договор + XML для ЭДО.\n"
            "Реквизиты покупателя: ИНН из кэша, текст, фото или файл.\n"
            "Продавец: ИП Шавкова Тамара Расуловна · НДС 5% (ОСНО)"
        )
        await bot.set_my_short_description(
            "Счёт + УПД + Договор + XML одним сообщением · Gemini AI"
        )
    except Exception:
        pass  # не критично

    dp = Dispatcher(storage=MemoryStorage())
    dp.include_router(router)

    await dp.start_polling(bot, allowed_updates=['message', 'callback_query'])


if __name__ == '__main__':
    asyncio.run(main())
